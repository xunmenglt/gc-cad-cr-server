# BSD 3- Clause License Copyright (c) 2023, Tecorigin Co., Ltd. All rights
# reserved.
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
# Redistributions of source code must retain the above copyright notice,
# this list of conditions and the following disclaimer.
# Redistributions in binary form must reproduce the above copyright notice,
# this list of conditions and the following disclaimer in the documentation
# and/or other materials provided with the distribution.
# Neither the name of the copyright holder nor the names of its contributors
# may be used to endorse or promote products derived from this software
# without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE
# ARE DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE
# LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR
# CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF
# SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS
# INTERRUPTION)
# HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT,
# STRICT LIABILITY,OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE)  ARISING IN ANY
# WAY OUT OF THE USE OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY
# OF SUCH DAMAGE.

from langchain_community.document_loaders import UnstructuredFileLoader
from typing import List
import tqdm

from docx.table import _Cell, Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx import Document, ImagePart
from PIL import Image
from io import BytesIO
from rag.module.indexing.loader.ocr import get_rapid_ocr
import numpy as np


class CustomizedOcrDocLoader(UnstructuredFileLoader):

    def _is_paragraph_end(self, text):
        if text.strip()[-1] in ["。", "？"]: return True
        else: return False

    def _get_elements(self) -> List:
        def doc2text(filepath):

            ocr = get_rapid_ocr()
            doc = Document(filepath)
            resp = ""

            def iter_block_items(parent):
                from docx.document import Document
                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("CustomizedOcrDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(total=len(doc.paragraphs)+len(doc.tables),
                               desc="CustomizedOcrDocLoader block index: 0")
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description(
                    "CustomizedOcrDocLoader block index: {}".format(i))
                b_unit.refresh()
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"

                    images = block._element.xpath('.//pic:pic')  # 获取所有图片
                    for image in images:
                        for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
                            part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
                            if isinstance(part, ImagePart):
                                image = Image.open(BytesIO(part._blob))
                                # result, _ = ocr(np.array(image))
                                result=None
                                if result:
                                    ocr_result = [line[1] for line in result]
                                    resp += "\n".join(ocr_result)

                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"

                b_unit.update(1)
            return resp

        text = doc2text(self.file_path)
        from unstructured.partition.text import partition_text
        return partition_text(text=text, **self.unstructured_kwargs)

